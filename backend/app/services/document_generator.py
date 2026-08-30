import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
import openpyxl
from openpyxl.utils import get_column_letter
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from django.utils import timezone
from app.models import DefenseCouncil, CouncilMember, GraduationProject, FinalGradeSummary, CouncilLiveScore

class DocumentGenerationService:

    @staticmethod
    def generate_to_trinh_docx(council_id):
        """
        Sinh file Word .docx Tờ trình thành lập Hội đồng chấm đồ án tốt nghiệp
        theo thể thức văn bản hành chính chuẩn của Trường ĐH Giao thông Vận tải (UTC).
        """
        council = DefenseCouncil.objects.select_related('batch').get(id=council_id)
        members = list(
            CouncilMember.objects.filter(council=council)
            .select_related('user', 'supervisor')
            .order_by('id')
        )

        doc = Document()

        # Set standard margins (1 inch)
        for section in doc.sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(0.8)

        # 1. Header Table (2 columns)
        header_table = doc.add_table(rows=1, cols=2)
        header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        header_table.autofit = False

        cell_left = header_table.cell(0, 0)
        cell_right = header_table.cell(0, 1)
        cell_left.width = Inches(3.2)
        cell_right.width = Inches(3.8)

        p_left = cell_left.paragraphs[0]
        p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_l1 = p_left.add_run("TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI\n")
        run_l1.font.name = "Times New Roman"
        run_l1.font.size = Pt(11)
        run_l1.font.bold = True
        run_l2 = p_left.add_run("KHOA CÔNG NGHỆ THÔNG TIN\n")
        run_l2.font.name = "Times New Roman"
        run_l2.font.size = Pt(11)
        run_l2.font.bold = True
        run_l3 = p_left.add_run("--------------------")
        run_l3.font.name = "Times New Roman"
        run_l3.font.size = Pt(9)

        p_right = cell_right.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_r1 = p_right.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n")
        run_r1.font.name = "Times New Roman"
        run_r1.font.size = Pt(11)
        run_r1.font.bold = True
        run_r2 = p_right.add_run("Độc lập – Tự do – Hạnh phúc\n")
        run_r2.font.name = "Times New Roman"
        run_r2.font.size = Pt(12)
        run_r2.font.bold = True
        run_r3 = p_right.add_run("--------------------")
        run_r3.font.name = "Times New Roman"
        run_r3.font.size = Pt(9)

        # Date
        today = timezone.now()
        p_date = doc.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_d = p_date.add_run(f"Hà Nội, ngày {today.day:02d} tháng {today.month:02d} năm {today.year}\n")
        run_d.font.name = "Times New Roman"
        run_d.font.size = Pt(12)
        run_d.font.italic = True

        # Title
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_t1 = p_title.add_run("TỜ TRÌNH\n")
        run_t1.font.name = "Times New Roman"
        run_t1.font.size = Pt(14)
        run_t1.font.bold = True
        run_t2 = p_title.add_run(f"V/v: Thành lập hội đồng chấm đồ án tốt nghiệp {council.council_name}\n")
        run_t2.font.name = "Times New Roman"
        run_t2.font.size = Pt(12)
        run_t2.font.bold = True

        # Kính gửi
        p_kinhgui = doc.add_paragraph()
        p_kinhgui.paragraph_format.left_indent = Inches(0.5)
        run_kg = p_kinhgui.add_run("Kính gửi:   - Ban Giám hiệu\n                 - Phòng Đào tạo Đại học\n                 - Khoa Đào tạo quốc tế\n")
        run_kg.font.name = "Times New Roman"
        run_kg.font.size = Pt(12)
        run_kg.font.bold = True

        # Body
        p_body = doc.add_paragraph()
        p_body.paragraph_format.first_line_indent = Inches(0.5)
        run_b1 = p_body.add_run(
            f"Căn cứ vào kế hoạch đào tạo năm học của Trường ĐH GTVT và tiến độ thực hiện đồ án tốt nghiệp đợt {council.batch.batch_name}, "
            f"Khoa Công nghệ Thông tin kính đề nghị Nhà trường cho phép thành lập Hội đồng chấm đồ án tốt nghiệp ({council.council_name}) "
            f"dành cho sinh viên hệ chính quy.\n\n"
            f"Đề xuất danh sách thành viên Hội đồng chấm đồ án tốt nghiệp gồm các Thầy/Cô sau:\n"
        )
        run_b1.font.name = "Times New Roman"
        run_b1.font.size = Pt(12)

        # Members Table
        table = doc.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table.rows[0].cells
        headers = ["STT", "Họ và tên", "Đơn vị công tác / Bộ môn", "Trách nhiệm trong HĐ"]
        col_widths = [Inches(0.6), Inches(2.2), Inches(2.2), Inches(2.0)]

        for i, text in enumerate(headers):
            hdr_cells[i].text = text
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in hdr_cells[i].paragraphs[0].runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(11)
                r.font.bold = True

        for idx, m in enumerate(members, start=1):
            row_cells = table.add_row().cells
            title = m.supervisor.academic_title if (m.supervisor and m.supervisor.academic_title) else ""
            full_name = f"{title} {m.user.get_full_name()}".strip()
            unit = m.external_institution if m.role == "EXTERNAL_MEMBER" and m.external_institution else (m.supervisor.department_name if m.supervisor else "Khoa CNTT")
            
            row_cells[0].text = str(idx)
            row_cells[1].text = full_name
            row_cells[2].text = unit or "Bộ môn CNTT"
            row_cells[3].text = m.get_role_display()

            for ci in range(4):
                p = row_cells[ci].paragraphs[0]
                if ci in [0, 3]:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(11)

        # Closing & Signature
        p_close = doc.add_paragraph()
        p_close.paragraph_format.first_line_indent = Inches(0.5)
        run_cl = p_close.add_run("\nTrân trọng cảm ơn sự quan tâm và phê duyệt của Ban Giám hiệu Nhà trường!\n")
        run_cl.font.name = "Times New Roman"
        run_cl.font.size = Pt(12)

        # Signature Table
        sig_table = doc.add_table(rows=1, cols=2)
        sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        sig_left = sig_table.cell(0, 0)
        sig_right = sig_table.cell(0, 1)
        sig_left.width = Inches(3.5)
        sig_right.width = Inches(3.5)

        p_sl = sig_left.paragraphs[0]
        run_sl = p_sl.add_run("Nơi nhận:\n- Như trên;\n- Lưu: VP Khoa CNTT.")
        run_sl.font.name = "Times New Roman"
        run_sl.font.size = Pt(10)
        run_sl.font.italic = True

        p_sr = sig_right.paragraphs[0]
        p_sr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sr1 = p_sr.add_run("TRƯỞNG KHOA CNTT\n\n\n\n\n")
        run_sr1.font.name = "Times New Roman"
        run_sr1.font.size = Pt(12)
        run_sr1.font.bold = True
        run_sr2 = p_sr.add_run("TS. Hoàng Văn Thông")
        run_sr2.font.name = "Times New Roman"
        run_sr2.font.size = Pt(12)
        run_sr2.font.bold = True

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    @staticmethod
    def generate_bien_ban_excel(council_id):
        """
        Sinh file Excel .xlsx Biên bản & Bảng tổng hợp điểm chấm bảo vệ
        Đồ án tốt nghiệp theo quy chế tín chỉ Trường ĐH Giao thông Vận tải.
        """
        council = DefenseCouncil.objects.select_related('batch').get(id=council_id)
        projects = list(GraduationProject.objects.filter(council=council).select_related('student__user', 'supervisor__user', 'reviewer__user').order_by('student__user__last_name', 'student__user__first_name'))

        wb = openpyxl.Workbook()
        ws = wb.active
        if ws is None:
            ws = wb.create_sheet()
        ws.title = f"Hội đồng {council.council_number}"

        # Styles
        font_title = Font(name="Times New Roman", size=14, bold=True)
        font_sub = Font(name="Times New Roman", size=11, italic=True)
        font_header = Font(name="Times New Roman", size=10, bold=True, color="FFFFFF")
        font_data = Font(name="Times New Roman", size=10)
        font_bold = Font(name="Times New Roman", size=10, bold=True)

        fill_header = PatternFill(start_color="003366", end_color="003366", fill_type="solid")
        fill_subtotal = PatternFill(start_color="F2F4F7", end_color="F2F4F7", fill_type="solid")
        align_center = Alignment(horizontal="center", vertical="center", wrap_text=True)
        align_left = Alignment(horizontal="left", vertical="center", wrap_text=True)
        align_right = Alignment(horizontal="right", vertical="center")

        thin_border = Border(
            left=Side(style='thin', color='CCCCCC'),
            right=Side(style='thin', color='CCCCCC'),
            top=Side(style='thin', color='CCCCCC'),
            bottom=Side(style='thin', color='CCCCCC')
        )

        # 1. Header rows
        ws.merge_cells('A1:G1')
        ws['A1'] = "TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI - KHOA CÔNG NGHỆ THÔNG TIN"
        ws['A1'].font = Font(name="Times New Roman", size=11, bold=True)

        ws.merge_cells('A3:N3')
        ws['A3'] = f"BẢNG TỔNG HỢP ĐIỂM CHẤM BẢO VỆ ĐỒ ÁN TỐT NGHIỆP - {council.council_name.upper()}"
        ws['A3'].font = font_title
        ws['A3'].alignment = align_center

        ws.merge_cells('A4:N4')
        ws['A4'] = f"Đợt đào tạo: {council.batch.batch_name} | Phòng bảo vệ: {council.defense_room or 'TBA'} | Thời gian: {council.session_date or 'TBA'}"
        ws['A4'].font = font_sub
        ws['A4'].alignment = align_center

        # 2. Table Headers
        headers = [
            ("STT", 6),
            ("Mã SV", 13),
            ("Họ và tên", 22),
            ("Lớp", 18),
            ("Tên đề tài đồ án", 35),
            ("GVHD", 20),
            ("GVPB", 20),
            ("Điểm GVHD\n(40%)", 12),
            ("Điểm GVPB\n(20%)", 12),
            ("Điểm HĐ\n(40%)", 12),
            ("Điểm Thang 10", 14),
            ("Điểm Thang 4", 13),
            ("Điểm Chữ", 11),
            ("Kết luận", 13)
        ]

        row_idx = 6
        for col_idx, (hdr_text, col_width) in enumerate(headers, start=1):
            cell = ws.cell(row=row_idx, column=col_idx, value=hdr_text)
            cell.font = font_header
            cell.fill = fill_header
            cell.alignment = align_center
            col_letter = get_column_letter(col_idx)
            ws.column_dimensions[col_letter].width = col_width

        # 3. Table Data
        for stt, proj in enumerate(projects, start=1):
            row_idx += 1
            student = proj.student
            user = student.user
            full_name = user.get_full_name()
            sup_name = proj.supervisor.user.get_full_name() if proj.supervisor else ""
            rev_name = proj.reviewer.user.get_full_name() if proj.reviewer else ""

            # Try get final grade summary
            summary = getattr(proj, 'final_grade_summary', None)
            if not summary:
                summary, _ = FinalGradeSummary.objects.get_or_create(project=proj)
                summary.supervisor_score = proj.supervisor_score
                summary.reviewer_score = proj.reviewer_score
                # average council score
                if hasattr(proj, 'council_scores'):
                    scores = list(proj.council_scores.all())
                else:
                    scores = list(CouncilLiveScore.objects.filter(project=proj))
                if scores:
                    summary.council_avg_score = round(sum(s.total_score for s in scores) / len(scores), 2)
                summary.calculate_and_save()

            row_data = [
                stt,
                student.registration_no,
                full_name,
                student.department or "CNTT",
                proj.topic_title_vi,
                sup_name,
                rev_name,
                summary.supervisor_score if summary.supervisor_score is not None else "",
                summary.reviewer_score if summary.reviewer_score is not None else "",
                summary.council_avg_score if summary.council_avg_score is not None else "",
                summary.final_score_10 if summary.final_score_10 is not None else "",
                summary.final_score_4 if summary.final_score_4 is not None else "",
                summary.final_letter_grade or "",
                "Đạt" if summary.is_passed else "Không đạt"
            ]

            for c_idx, val in enumerate(row_data, start=1):
                c = ws.cell(row=row_idx, column=c_idx, value=val)
                c.font = font_data
                c.border = thin_border
                if c_idx in [1, 2, 8, 9, 10, 11, 12, 13, 14]:
                    c.alignment = align_center
                elif c_idx in [3, 4, 5, 6, 7]:
                    c.alignment = align_left

        # Save buffer
        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return buffer
