import io
import openpyxl
from docx import Document
from django.test import TestCase
from django.db import IntegrityError
from rest_framework.test import APIClient
from rest_framework import status

from app.models import (
    CustomUser,
    AcademicBatch,
    CourseClass,
    Student,
    Supervisor,
    SupervisorQuota,
    ProjectTopicArea,
    InternshipInfo,
    GraduationProject,
    WeeklyProgressReport,
    DefenseCouncil,
    CouncilMember,
    CouncilLiveScore,
    EvaluationPolicy,
    FinalGradeSummary,
    AuditLog
)
from app.services.excel_importer import ExcelImportService
from app.services.allocation_engine import MinCostMaxFlowAllocationEngine
from app.services.reviewer_engine import ReviewerAndCouncilAllocationEngine
from app.services.document_generator import DocumentGenerationService

class AuditDeepVerificationTests(TestCase):
    def setUp(self):
        self.client = APIClient()

        # Admin user
        self.admin = CustomUser.objects.create_user(
            username="admin_audit",
            email="admin_audit@utc.edu.vn",
            password="adminpassword123",
            user_type="admin",
            is_staff=True
        )

        # Batch & Policy
        self.batch = AcademicBatch.objects.create(
            batch_code="2026_2027_HK1_AUDIT",
            batch_name="Đợt ĐATN Kiểm toán Toàn diện",
            is_active=True
        )
        self.policy = EvaluationPolicy.objects.create(
            batch=self.batch,
            weight_supervisor=0.4,
            weight_reviewer=0.2,
            weight_council=0.4
        )

        # Topic Area
        self.topic_software = ProjectTopicArea.objects.create(
            name="Phát triển phần mềm và ứng dụng (webApp, MobileApp)",
            code="SOFTWARE_DEV",
            is_active=True
        )
        self.topic_ai = ProjectTopicArea.objects.create(
            name="Dữ liệu và trí tuệ nhân tạo",
            code="AI_DATA",
            is_active=True
        )

        # 3 Faculty members
        self.sup1_user = CustomUser.objects.create_user(
            username="gv_du_audit", email="du@utc.edu.vn", password="password123",
            first_name="Dư", last_name="Nguyễn Đức", user_type="supervisor"
        )
        self.sup1 = Supervisor.objects.create(
            user=self.sup1_user, supervisor_id="GV003", academic_title="TS", department_name="CNPM"
        )
        self.quota1 = SupervisorQuota.objects.create(
            supervisor=self.sup1, batch=self.batch, viet_anh_quota=2, general_cntt_quota=3, max_total_quota=5
        )

        self.sup2_user = CustomUser.objects.create_user(
            username="gv_sao_audit", email="sao@utc.edu.vn", password="password123",
            first_name="Sao", last_name="Nguyễn Kim", user_type="supervisor"
        )
        self.sup2 = Supervisor.objects.create(
            user=self.sup2_user, supervisor_id="GV008", academic_title="TS", department_name="Mạng&HTTT"
        )
        self.quota2 = SupervisorQuota.objects.create(
            supervisor=self.sup2, batch=self.batch, viet_anh_quota=2, general_cntt_quota=2, max_total_quota=4
        )

        self.sup3_user = CustomUser.objects.create_user(
            username="gv_long_audit", email="long@utc.edu.vn", password="password123",
            first_name="Long", last_name="Nguyễn Văn", user_type="supervisor"
        )
        self.sup3 = Supervisor.objects.create(
            user=self.sup3_user, supervisor_id="GV020", academic_title="PGS.TS", department_name="KHMT"
        )
        self.quota3 = SupervisorQuota.objects.create(
            supervisor=self.sup3, batch=self.batch, viet_anh_quota=0, general_cntt_quota=3, max_total_quota=3
        )

    # =========================================================================
    # PHẦN 5: KIỂM TRA THUẬT TOÁN MCMF (Scenarios A, B, C, D, E, F)
    # =========================================================================

    def test_mcmf_scenario_a_sufficient_quota(self):
        """Scenario A: Sufficient quota, all students allocated to top preferences"""
        for i in range(3):
            u = CustomUser.objects.create_user(username=f"sv_a_{i}", user_type="student")
            s = Student.objects.create(user=u, registration_no=f"2112000A{i}", academic_batch=self.batch)
            InternshipInfo.objects.create(
                student=s, batch=self.batch, preferred_supervisor=self.sup1, topic_direction=self.topic_software
            )

        res = MinCostMaxFlowAllocationEngine.allocate_supervisors_for_batch(self.batch.id)
        self.assertTrue(res["success"])
        self.assertEqual(res["matched_count"], 3)
        self.assertEqual(res["unassigned_count"], 0)
        # All 3 matched to Sup1 since Sup1 quota is 5
        projects = GraduationProject.objects.filter(batch=self.batch, supervisor=self.sup1)
        self.assertEqual(projects.count(), 3)

    def test_mcmf_scenario_b_insufficient_quota(self):
        """Scenario B: Total student count exceeds total faculty quota capacity"""
        # Create 15 students when total quota is 5 + 4 + 3 = 12
        for i in range(15):
            u = CustomUser.objects.create_user(username=f"sv_b_{i}", user_type="student")
            s = Student.objects.create(user=u, registration_no=f"2112000B{i}", academic_batch=self.batch)
            InternshipInfo.objects.create(
                student=s, batch=self.batch, topic_direction=self.topic_software
            )

        res = MinCostMaxFlowAllocationEngine.allocate_supervisors_for_batch(self.batch.id)
        self.assertTrue(res["success"])
        self.assertEqual(res["matched_count"], 12)  # Exactly total capacity
        self.assertEqual(res["unassigned_count"], 3)  # 3 remain unassigned with clear report

        # Verify each supervisor does not exceed max_total_quota
        self.assertEqual(GraduationProject.objects.filter(batch=self.batch, supervisor=self.sup1).count(), 5)
        self.assertEqual(GraduationProject.objects.filter(batch=self.batch, supervisor=self.sup2).count(), 4)
        self.assertEqual(GraduationProject.objects.filter(batch=self.batch, supervisor=self.sup3).count(), 3)

    def test_mcmf_scenario_c_multiple_students_same_supervisor(self):
        """Scenario C: 10 students pick Sup1 (Quota=5) -> 5 get Sup1, others matched to next best affinity"""
        for i in range(10):
            u = CustomUser.objects.create_user(username=f"sv_c_{i}", user_type="student")
            s = Student.objects.create(user=u, registration_no=f"2112000C{i}", academic_batch=self.batch)
            InternshipInfo.objects.create(
                student=s, batch=self.batch, preferred_supervisor=self.sup1, topic_direction=self.topic_software
            )

        res = MinCostMaxFlowAllocationEngine.allocate_supervisors_for_batch(self.batch.id)
        self.assertTrue(res["success"])
        self.assertEqual(res["matched_count"], 10)
        # Exactly 5 matched to Sup1
        self.assertEqual(GraduationProject.objects.filter(batch=self.batch, supervisor=self.sup1).count(), 5)

    def test_mcmf_scenario_d_supervisor_full_quota(self):
        """Scenario D: Supervisor with quota 0 cannot receive any student"""
        self.quota1.max_total_quota = 0
        self.quota1.save()

        u = CustomUser.objects.create_user(username="sv_d_1", user_type="student")
        s = Student.objects.create(user=u, registration_no="2112000D1", academic_batch=self.batch)
        InternshipInfo.objects.create(student=s, batch=self.batch, preferred_supervisor=self.sup1)

        res = MinCostMaxFlowAllocationEngine.allocate_supervisors_for_batch(self.batch.id)
        self.assertTrue(res["success"])
        # Sup1 receives 0 students
        self.assertEqual(GraduationProject.objects.filter(batch=self.batch, supervisor=self.sup1).count(), 0)

    def test_mcmf_scenario_e_no_supervisors_in_batch(self):
        """Scenario E: Error report when no quotas exist for batch"""
        new_batch = AcademicBatch.objects.create(batch_code="EMPTY_BATCH")
        u = CustomUser.objects.create_user(username="sv_e_1", user_type="student")
        Student.objects.create(user=u, registration_no="2112000E1", academic_batch=new_batch)
        res = MinCostMaxFlowAllocationEngine.allocate_supervisors_for_batch(new_batch.id)
        self.assertFalse(res["success"])
        self.assertIn("Chưa thiết lập định mức Quota", str(res["error"]))

    def test_mcmf_scenario_f_second_run_idempotency(self):
        """Scenario F: Running auto-match a 2nd time preserves project IDs and updates quota stats cleanly"""
        u = CustomUser.objects.create_user(username="sv_f_1", user_type="student")
        s = Student.objects.create(user=u, registration_no="2112000F1", academic_batch=self.batch)
        InternshipInfo.objects.create(student=s, batch=self.batch, preferred_supervisor=self.sup1)

        res1 = MinCostMaxFlowAllocationEngine.allocate_supervisors_for_batch(self.batch.id)
        proj1 = GraduationProject.objects.get(student=s)

        res2 = MinCostMaxFlowAllocationEngine.allocate_supervisors_for_batch(self.batch.id)
        proj2 = GraduationProject.objects.get(student=s)

        self.assertEqual(proj1.id, proj2.id)
        self.assertEqual(res2["matched_count"], 1)

    # =========================================================================
    # PHẦN 6: KIỂM TRA NO-CONFLICT CONSTRAINT
    # =========================================================================

    def test_reviewer_allocation_strict_no_conflict(self):
        """Reviewer MUST NOT be Supervisor and Supervisor MUST NOT be in Council"""
        # Create Project with Sup1
        u = CustomUser.objects.create_user(username="sv_nc_1", user_type="student")
        s = Student.objects.create(user=u, registration_no="2112000NC1", academic_batch=self.batch)
        p = GraduationProject.objects.create(
            student=s, supervisor=self.sup1, batch=self.batch, topic_title_vi="Đề tài An toàn No-Conflict"
        )

        # Council 1 contains Sup1 (Chair) and Sup2 (Secretary)
        council1 = DefenseCouncil.objects.create(batch=self.batch, council_number=1, council_name="HĐ 1 (Có Sup1)")
        CouncilMember.objects.create(council=council1, supervisor=self.sup1, user=self.sup1_user, role="CHAIR")
        CouncilMember.objects.create(council=council1, supervisor=self.sup2, user=self.sup2_user, role="SECRETARY")

        # Council 2 contains Sup2 (Chair) and Sup3 (Member) (Does NOT contain Sup1)
        council2 = DefenseCouncil.objects.create(batch=self.batch, council_number=2, council_name="HĐ 2 (Không có Sup1)")
        CouncilMember.objects.create(council=council2, supervisor=self.sup2, user=self.sup2_user, role="CHAIR")
        CouncilMember.objects.create(council=council2, supervisor=self.sup3, user=self.sup3_user, role="MEMBER")

        res = ReviewerAndCouncilAllocationEngine.assign_councils_and_reviewers(self.batch.id)
        self.assertTrue(res["success"])

        p.refresh_from_db()
        # MUST be assigned to Council 2 (since Sup1 is in Council 1)
        self.assertEqual(p.council, council2)
        # Reviewer must be Sup2 or Sup3 (NOT Sup1)
        self.assertIn(p.reviewer, [self.sup2, self.sup3])
        self.assertNotEqual(p.reviewer, p.supervisor)

    # =========================================================================
    # PHẦN 9: KIỂM TRA EXCEL IMPORT (IDEMPOTENCY & DUPLICATES)
    # =========================================================================

    def test_excel_import_idempotency_and_duplicate_prevention(self):
        """Importing the same roster twice does not duplicate students/users"""
        wb = openpyxl.Workbook()
        ws = wb.active
        assert ws is not None
        ws.title = "IT1.659.103"
        ws.cell(row=6, column=1, value="Học phần: Đồ án tốt nghiệp Kỹ sư CNTT")
        ws.cell(row=9, column=1, value="STT")
        ws.cell(row=9, column=2, value="Lớp")
        ws.cell(row=9, column=3, value="Mã số SV")
        ws.cell(row=9, column=4, value="Họ đệm")
        ws.cell(row=9, column=5, value="Tên")

        # Row 1 with Vietnamese Unicode
        ws.cell(row=10, column=1, value=1)
        ws.cell(row=10, column=2, value="Kỹ sư CNTT 1 K62")
        ws.cell(row=10, column=3, value="211299999")
        ws.cell(row=10, column=4, value="Đặng Hoàng")
        ws.cell(row=10, column=5, value="Dương")

        buf1 = io.BytesIO()
        wb.save(buf1)
        buf1.seek(0)

        # First import
        res1 = ExcelImportService.import_students_from_excel(buf1, self.batch.id)
        self.assertEqual(res1["created"], 1)

        # Second import with same file
        buf2 = io.BytesIO()
        wb.save(buf2)
        buf2.seek(0)
        res2 = ExcelImportService.import_students_from_excel(buf2, self.batch.id)
        self.assertEqual(res2["created"], 0)  # 0 new students created
        self.assertEqual(res2["updated"], 1)  # 1 updated cleanly

        # Verify exactly 1 student in DB
        self.assertEqual(Student.objects.filter(registration_no="211299999").count(), 1)

    # =========================================================================
    # PHẦN 10: KIỂM TRA DOCUMENT GENERATION (WORD & EXCEL)
    # =========================================================================

    def test_document_generation_word_and_excel_complete(self):
        """Verify dynamic Word Tờ trình and Excel Bảng điểm"""
        council = DefenseCouncil.objects.create(
            batch=self.batch, council_number=1, council_name="Hội đồng 1 - CNPM",
            defense_room="502-A9", session_date="2026-06-15"
        )
        CouncilMember.objects.create(council=council, supervisor=self.sup1, user=self.sup1_user, role="CHAIR")
        CouncilMember.objects.create(council=council, supervisor=self.sup2, user=self.sup2_user, role="SECRETARY")

        # Word Document
        docx_buf = DocumentGenerationService.generate_to_trinh_docx(council.id)
        doc = Document(docx_buf)
        paragraphs_text = " ".join([p.text for p in doc.paragraphs])
        tables_text = " ".join([cell.text for t in doc.tables for row in t.rows for cell in row.cells])
        full_text = paragraphs_text + " " + tables_text

        self.assertIn("TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI", full_text)
        self.assertIn("KHOA CÔNG NGHỆ THÔNG TIN", full_text)
        self.assertIn("TỜ TRÌNH", full_text)
        self.assertIn("Hội đồng 1 - CNPM", full_text)
        self.assertIn("TS. Hoàng Văn Thông", full_text)

        # Excel Document
        xlsx_buf = DocumentGenerationService.generate_bien_ban_excel(council.id)
        wb = openpyxl.load_workbook(xlsx_buf)
        self.assertIn("Hội đồng 1", wb.sheetnames)

    # =========================================================================
    # PHẦN 11: FULL 20-STEP END-TO-END WORKFLOW INTEGRATION TEST
    # =========================================================================

    def test_full_20_step_e2e_lifecycle(self):
        """
        Complete 20-Step Lifecycle:
        1. Batch -> 2. Import SV -> 3. Load Supervisors -> 4. Quotas -> 5. Survey ->
        6. MCMF -> 7. Project Created -> 8. Outline -> 9. Review -> 10. Weekly Reports ->
        11. Supervisor Evaluation -> 12. Council -> 13. Members -> 14. Reviewer ->
        15. Conflict Check -> 16. Defense Session -> 17. Live Grading -> 18. Final Grade (40/20/40) ->
        19. Export Word -> 20. Export Excel
        """
        # Step 1: Batch & Policy ready in setUp
        # Step 2: Create Student
        u = CustomUser.objects.create_user(username="211299901", user_type="student", email="duong@utc.edu.vn")
        st = Student.objects.create(user=u, registration_no="211299901", academic_batch=self.batch)

        # Step 5: Student submits preference
        InternshipInfo.objects.create(
            student=st, batch=self.batch, is_interning=True, company_name="FPT Software",
            preferred_supervisor=self.sup1, topic_direction=self.topic_software
        )

        # Step 6 & 7: Auto match MCMF
        res_mcmf = MinCostMaxFlowAllocationEngine.allocate_supervisors_for_batch(self.batch.id)
        self.assertTrue(res_mcmf["success"])

        proj = GraduationProject.objects.get(student=st)
        self.assertEqual(proj.supervisor, self.sup1)
        self.assertEqual(proj.status, "ALLOCATED")

        # Step 8 & 9: Outline submission & review
        proj.status = "OUTLINE_APPROVED"
        proj.save()

        # Step 10: Weekly Progress
        for w in range(1, 16):
            WeeklyProgressReport.objects.create(
                project=proj, week_number=w, summary_content=f"Hoàn thành công việc tuần {w}",
                supervisor_rating="GOOD", supervisor_feedback="Tiến độ tốt"
            )
        self.assertEqual(WeeklyProgressReport.objects.filter(project=proj).count(), 15)

        # Step 11: Supervisor evaluates (Score = 8.0)
        proj.supervisor_score = 8.0
        proj.is_eligible_for_defense = True
        proj.status = "DEFENSE_READY"
        proj.save()

        # Step 12 & 13: Create Council & Members (Sup2 & Sup3, NOT Sup1)
        council = DefenseCouncil.objects.create(batch=self.batch, council_number=1, council_name="HĐ 1 - K62")
        m2 = CouncilMember.objects.create(council=council, supervisor=self.sup2, user=self.sup2_user, role="CHAIR")
        m3 = CouncilMember.objects.create(council=council, supervisor=self.sup3, user=self.sup3_user, role="SECRETARY")

        # Step 14 & 15: Assign Reviewer & No Conflict
        res_rev = ReviewerAndCouncilAllocationEngine.assign_councils_and_reviewers(self.batch.id)
        self.assertTrue(res_rev["success"])
        proj.refresh_from_db()
        self.assertEqual(proj.council, council)
        self.assertEqual(proj.reviewer, self.sup2)

        # Step 16 & 17: Reviewer score (7.0) & Council Live Score (9.0)
        proj.reviewer_score = 7.0
        proj.save()

        # Council Chair m2 gives 9.0 (Presentation=2.7, Content=2.7, QA=1.8, Demo=1.8)
        score_obj = CouncilLiveScore.objects.create(
            council=council, project=proj, member=m2,
            score_presentation=2.7, score_content=2.7, score_qa=1.8, score_demo=1.8
        )
        self.assertEqual(score_obj.total_score, 9.0)

        # Step 18: Calculate Final Grade (40% GVHD=8.0 + 20% GVPB=7.0 + 40% HĐ=9.0)
        # Expected: 8.0*0.4 + 7.0*0.2 + 9.0*0.4 = 3.2 + 1.4 + 3.6 = 8.2 (B+, Khá, ĐẠT)
        summary, _ = FinalGradeSummary.objects.get_or_create(project=proj)
        summary.supervisor_score = 8.0
        summary.reviewer_score = 7.0
        summary.council_avg_score = 9.0
        summary.calculate_and_save(policy=self.policy)

        self.assertEqual(summary.final_score_10, 8.2)
        self.assertEqual(summary.final_score_4, 3.5)
        self.assertEqual(summary.final_letter_grade, "B+")
        self.assertEqual(summary.classification, "Khá giỏi")
        self.assertTrue(summary.is_passed)

        # Step 19 & 20: Export Word & Excel
        doc_buf = DocumentGenerationService.generate_to_trinh_docx(council.id)
        self.assertGreater(len(doc_buf.getvalue()), 1000)

        xls_buf = DocumentGenerationService.generate_bien_ban_excel(council.id)
        self.assertGreater(len(xls_buf.getvalue()), 1000)
